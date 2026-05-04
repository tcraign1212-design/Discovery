import streamlit as st
import os
import io
import re
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
import google.generativeai as genai
from openai import OpenAI

# 1. Page Configuration
st.set_page_config(page_title="Discovery Drafter & Utility", layout="wide")

st.title("Legal Utility: PDF Converter & Discovery Response Drafter")
st.markdown("---")

# Initialize taxonomy dictionary
TAXONOMY_OBJECTIONS = dict()

# Standalone Keys
K1 = "OB-TX-001: Relevance / Outside Scope"
K2 = "OB-TX-002: Overbroad as to Time"
K3 = "OB-TX-003: Overbroad as to Subject Matter"
K4 = "OB-TX-004: Vague / Ambiguous / Undefined Terms"
K5 = "OB-TX-005: Lack of Reasonable Particularity"
K6 = "OB-TX-006: Improper Fishing Expedition"
K7 = "OB-TX-007: Undue Burden / Expense"
K8 = "OB-TX-008: More Convenient / Less Burdensome Source"
K9 = "OB-TX-009: Duplicative / Previously Produced"
K10 = "OB-TX-010: Premature / Discovery Ongoing"
K11 = "OB-TX-011: Mandatory Initial Disclosures (TRCP 194)"
K12 = "OB-TX-012: TRCP 193.3 Withholding Statement"
K13 = "OB-TX-013: Attorney-Client Privilege"
K14 = "OB-TX-014: Work Product Privilege"
K15 = "OB-TX-015: Consulting Expert Protection"
K16 = "OB-TX-016: Spousal Privilege"
K17 = "OB-TX-017: Mental Health Records Not at Issue"
K18 = "OB-TX-018: Privacy / Sensitive Personal Info"
K19 = "OB-TX-019: Narrative / Marshaling Proof"
K20 = "OB-TX-020: Marshal All Evidence"
K21 = "OB-TX-021: Medical Opinion from Lay Party"
K22 = "OB-TX-022: Improper Expert Discovery by Interrogatory"
K23 = "OB-TX-023: Beyond Current Knowledge"
K24 = "OB-TX-024: Exceeds Numerical Discovery Plan Limits"
K25 = "OB-TX-025: Not in Possession, Custody, or Control"
K26 = "OB-TX-026: Request Requires Creation of a Document"
K27 = "OB-TX-027: Blank Authorization / Lack of Specificity"
K28 = "OB-TX-028: Medical Authorization Improper / Records Forthcoming"
K29 = "OB-TX-029: Tax Returns / Heightened Financial Privacy"
K30 = "OB-TX-030: Social Security / Identifier Forms"
K31 = "OB-TX-031: Employment Records Limited to Wage Period"
K32 = "OB-TX-032: Phone Records Lack Nexus"
K33 = "OB-TX-033: Collateral Source Rule"
K34 = "OB-TX-034: Medical Billing Creation Expansion"
K35 = "OB-TX-035: Premature Damage Computation"
K36 = "OB-TX-036: Core Issue / Disputed Merits RFA"
K37 = "OB-TX-037: Compound RFA"
K38 = "OB-TX-038: Calls for Legal Conclusion"
K39 = "OB-TX-039: Cannot Admit/Deny After Reasonable Inquiry"
K40 = "OB-TX-040: Prejudicial Terms / Controverted Text"
K41 = "OB-TX-041: All Written Complaints (Overbroad)"
K42 = "OB-TX-042: Ex Parte Provider Communications"

# Mapping Values
V1 = "Plaintiff objects to this request under TRCP 193.2 to the extent it seeks information not relevant to any party's claim or defense and is not within the permissible scope of discovery."
V2 = "Plaintiff objects that this request is overly broad because it is not reasonably limited in time to the matters at issue in this litigation."
V3 = "Plaintiff objects that this request is overly broad because it is not reasonably tailored to include only matters relevant to the issues in dispute."
V4 = "Plaintiff objects that this request is vague and ambiguous because it fails to define the key terms with reasonable certainty, such that Plaintiff cannot determine the exact information sought."
V5 = "Plaintiff objects that this request fails to describe the items or documents sought with reasonable particularity and therefore imposes an improper burden on the responding party."
V6 = "Plaintiff objects that this request constitutes an improper fishing expedition and is not reasonably tailored to obtain information directly relevant to the claims or defenses at issue."
V7 = "Plaintiff objects that complying with this request would impose an undue burden and expense that is completely disproportionate to the likely benefit of the discovery."
V8 = "Plaintiff objects to this request to the extent the information or responsive documents sought are obtainable from a source that is more convenient, less burdensome, or less expensive."
V9 = "Plaintiff objects that this request is unreasonably cumulative or duplicative and, to the extent responsive material exists, it has already been produced or identified in prior productions."
V10 = "Plaintiff objects that this request is premature to the extent it seeks a complete factual or evidentiary statement before discovery is sufficiently developed."
V11 = "Plaintiff objects that this request seeks information through an improper discovery vehicle as the subject matter is directly governed by required initial disclosures under TRCP 194.1."
V12 = "Responsive material has been withheld pursuant to Tex. R. Civ. P. 193.3. The withheld material is responsive to this request and is withheld on the basis of applicable privileges."
V13 = "Plaintiff withholds responsive material protected by the attorney-client privilege and provides this withholding statement pursuant to Rule 193.3."
V14 = "Plaintiff objects and withholds responsive material to the extent it constitutes work product or material prepared in anticipation of litigation under TRCP 192.5."
V15 = "Plaintiff objects to this request to the extent it seeks the identity, mental impressions, or opinions of consulting experts not expected to testify."
V16 = "Plaintiff objects to the extent this request seeks confidential communications between spouses protected by the spousal communications privilege."
V17 = "Plaintiff objects to this request to the extent it seeks highly sensitive mental-health information where Plaintiff has not affirmatively placed their mental condition at issue in this litigation."
V18 = "Plaintiff objects to this request to the extent it seeks highly sensitive personal identifiers or private information without a demonstrated need that is proportional to the case."
V19 = "Plaintiff objects to this interrogatory to the extent it requires a narrative response or detailed marshaling of proof more appropriately developed through deposition."
V20 = "Plaintiff objects that this interrogatory improperly seeks to force Plaintiff to marshal all evidence supporting its claims or defenses."
V21 = "Plaintiff objects to the extent this interrogatory requires a lay party Plaintiff to provide medical opinions beyond Plaintiff's personal knowledge or expert qualifications."
V22 = "Plaintiff objects to this interrogatory to the extent it seeks expert information outside the scope or manner authorized by the expert discovery rules."
V23 = "Plaintiff objects to the extent this interrogatory seeks information beyond Plaintiff's present knowledge and attempts to bind Plaintiff to a complete evidentiary statement before discovery is complete."
V24 = "Plaintiff objects because this set exceeds the maximum number of permissible requests or answers under the governing TRCP discovery control plan."
V25 = "Plaintiff objects to the extent this request seeks materials not within Plaintiff's possession, custody, or control."
V26 = "Plaintiff objects because this request improperly requires Plaintiff to create a document that does not presently exist."
V27 = "Plaintiff objects to signing the requested authorization in blank because it fails to specify the records sought or specific providers, depriving Plaintiff of a meaningful opportunity to evaluate relevance."
V28 = "Plaintiff objects to the extent this request seeks a blanket medical authorization rather than relevant medical records properly subject to production or disclosure."
V29 = "Plaintiff objects that this request seeks highly confidential tax information and is overbroad, intrusive, and not shown to be necessary in the form requested."
V30 = "Plaintiff objects to this request to the extent it seeks Social Security records or identifiers that are irrelevant, overbroad, and unduly intrusive."
V31 = "Plaintiff objects to this request to the extent it seeks personnel and employment records beyond those reasonably related to any claimed wage loss."
V32 = "Plaintiff objects because this request seeks telephone records without a pleaded or otherwise demonstrated nexus to any claim or defense in the case."
V33 = "Plaintiff objects to this request to the extent it seeks information regarding non-utilized health insurance, private health plans, or collateral benefits barred by the Texas Collateral Source Rule."
V34 = "Plaintiff objects to the extent this request improperly expands the burden of medical billing disclosure by requiring the creation of an itemized analysis beyond the records themselves."
V35 = "Plaintiff objects to the extent this request seeks a premature, exhaustive, or artificially fixed statement of damages before discovery is fully developed."
V36 = "Plaintiff objects because this request for admission improperly seeks to establish a disputed merits issue rather than narrow an uncontroverted fact."
V37 = "Plaintiff objects because this request for admission is compound and does not permit a fair admission or denial of a single proposition."
V38 = "Plaintiff objects to this request for admission to the extent it seeks a pure legal conclusion."
V39 = "After reasonable inquiry, Plaintiff lacks sufficient information to admit or deny this request and therefore denies it."
V40 = "Plaintiff objects to the use of highly subjective or prejudicial terms within the request to the extent they assume a disputed characterization of the incident."
V41 = "Plaintiff objects because the request for all written complaints across unrelated matters is overly broad, vague, and constitutes a prohibited fishing expedition."
V42 = "Plaintiff objects to the extent the requested authorization would permit ex parte communications with healthcare providers rather than the production of defined records."

TAXONOMY_OBJECTIONS[K1] = V1
TAXONOMY_OBJECTIONS[K2] = V2
TAXONOMY_OBJECTIONS[K3] = V3
TAXONOMY_OBJECTIONS[K4] = V4
TAXONOMY_OBJECTIONS[K5] = V5
TAXONOMY_OBJECTIONS[K6] = V6
TAXONOMY_OBJECTIONS[K7] = V7
TAXONOMY_OBJECTIONS[K8] = V8
TAXONOMY_OBJECTIONS[K9] = V9
TAXONOMY_OBJECTIONS[K10] = V10
TAXONOMY_OBJECTIONS[K11] = V11
TAXONOMY_OBJECTIONS[K12] = V12
TAXONOMY_OBJECTIONS[K13] = V13
TAXONOMY_OBJECTIONS[K14] = V14
TAXONOMY_OBJECTIONS[K15] = V15
TAXONOMY_OBJECTIONS[K16] = V16
TAXONOMY_OBJECTIONS[K17] = V17
TAXONOMY_OBJECTIONS[K18] = V18
TAXONOMY_OBJECTIONS[K19] = V19
TAXONOMY_OBJECTIONS[K20] = V20
TAXONOMY_OBJECTIONS[K21] = V21
TAXONOMY_OBJECTIONS[K22] = V22
TAXONOMY_OBJECTIONS[K23] = V23
TAXONOMY_OBJECTIONS[K24] = V24
TAXONOMY_OBJECTIONS[K25] = V25
TAXONOMY_OBJECTIONS[K26] = V26
TAXONOMY_OBJECTIONS[K27] = V27
TAXONOMY_OBJECTIONS[K28] = V28
TAXONOMY_OBJECTIONS[K29] = V29
TAXONOMY_OBJECTIONS[K30] = V30
TAXONOMY_OBJECTIONS[K31] = V31
TAXONOMY_OBJECTIONS[K32] = V32
TAXONOMY_OBJECTIONS[K33] = V33
TAXONOMY_OBJECTIONS[K34] = V34
TAXONOMY_OBJECTIONS[K35] = V35
TAXONOMY_OBJECTIONS[K36] = V36
TAXONOMY_OBJECTIONS[K37] = V37
TAXONOMY_OBJECTIONS[K38] = V38
TAXONOMY_OBJECTIONS[K39] = V39
TAXONOMY_OBJECTIONS[K40] = V40
TAXONOMY_OBJECTIONS[K41] = V41
TAXONOMY_OBJECTIONS[K42] = V42


# --- DOCUMENT GENERATION FUNCTIONS ---

def convert_pdf_text_to_docx(pdf_file):
    reader = PdfReader(pdf_file)
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            lines = page_text.split('\n')
            for line in lines:
                cleaned_line = line.strip()
                if cleaned_line:
                    doc.add_paragraph(cleaned_line)
        else:
            doc.add_paragraph(f"[Page {i+1} contained no directly extractable text]")
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_response_docx(content_text, include_styling_box=False, style_details=None):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    if include_styling_box and style_details:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        col_widths = [Inches(3.25), Inches(3.25)]
        for i, col in enumerate(table.columns):
            col.width = col_widths[i]

        cell_left = table.cell(0, 0)
        cell_right = table.cell(0, 1)

        p_left = cell_left.paragraphs[0]
        p_left.paragraph_format.space_after = Pt(0)
        p_left.paragraph_format.line_spacing = 1.15
        p_left.add_run(f"CAUSE NO. {style_details.get('cause_no', '')}\n").bold = True
        p_left.add_run(f"{style_details.get('plaintiff', '')}\nPlaintiff,\nvs.\n{style_details.get('defendant', '')}\nDefendant.")

        p_right = cell_right.paragraphs[0]
        p_right.paragraph_format.space_after = Pt(0)
        p_right.paragraph_format.line_spacing = 1.15
        p_right.add_run(f"§\tIN THE DISTRICT COURT\n§\n§\n§\t{style_details.get('court', '')}\n§\n§\t{style_details.get('county', '')} COUNTY, TEXAS")
        
        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_before = Pt(12)
        p_spacer.paragraph_format.space_after = Pt(12)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(18)
    p_title.paragraph_format.space_before = Pt(12)
    run_title = p_title.add_run("PLAINTIFF’S RESPONSES AND OBJECTIONS TO DEFENDANT’S DISCOVERY")
    run_title.font.bold = True
    run_title.font.size = Pt(12)

    # Patterns matching subparts: e.g., 'a. ', 'b. ', '(a) ', '(b) '
    subpart_regex = re.compile(r'^(?:[a-g1-9]\.|\([a-g1-9]\))\s')

    lines = content_text.split('\n')
    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line:
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15
        
        # Strip internal markers if passed down
        if cleaned_line.startswith('#'):
            cleaned_line = cleaned_line.replace('#', '').strip()

        upper_line = cleaned_line.upper()

        # Target absolute formatting matches
        is_discovery_item = False
        is_answer_item = False

        # Specific keyword matching
        if "INTERROGATORY NO." in upper_line:
            is_discovery_item = True
        elif "REQUEST FOR PRODUCTION NO." in upper_line:
            is_discovery_item = True
        elif "REQUEST FOR ADMISSION NO." in upper_line:
            is_discovery_item = True
        elif upper_line.startswith("ANSWER:") or upper_line.startswith("RESPONSE:"):
            is_answer_item = True

        if is_discovery_item:
            # Parse prefix vs rest of the query
            parts = re.split(r'(?i)(INTERROGATORY NO\.\s*\d+:|REQUEST FOR PRODUCTION NO\.\s*\d+:|REQUEST FOR ADMISSION NO\.\s*\d+:)', cleaned_line)
            for part in parts:
                p_part = part.strip()
                if not p_part:
                    continue
                if any(x in p_part.upper() for x in ["INTERROGATORY", "REQUEST FOR PRODUCTION", "REQUEST FOR ADMISSION"]):
                    run = p.add_run(part + " ")
                    run.font.bold = True
                else:
                    p.add_run(part)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)

        elif is_answer_item:
            run = p.add_run(cleaned_line)
            run.font.bold = True
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        elif subpart_regex.match(cleaned_line):
            # Indent exactly .5" per request
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            p.add_run(cleaned_line)

        else:
            # Standard output format fallback
            if "**" in cleaned_line:
                parts = cleaned_line.split('**')
                for index, part in enumerate(parts):
                    if index % 2 == 1:
                        p.add_run(part).font.bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(cleaned_line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# 3. Streamlit Tab UI
tab1, tab2 = st.tabs(["📄 Convert PDF to Word", "⚖️ Discovery Response Drafter"])

# --- TAB 1: PDF TO WORD ---
with tab1:
    st.subheader("Direct PDF to Word Converter")
    st.info("Upload any plain text PDF to convert it directly into an editable Word (.docx) file.")
    
    uploaded_pdf = st.file_uploader("Upload PDF File", type=["pdf"])
    
    if uploaded_pdf:
        with st.spinner("Parsing PDF text and building Word document..."):
            docx_buffer = convert_pdf_text_to_docx(uploaded_pdf)
            st.success("Conversion successful!")
            
            base_name = os.path.splitext(uploaded_pdf.name)[0]
            st.download_button(
                label=f"📥 Download '{base_name}.docx'",
                data=docx_buffer,
                file_name=f"{base_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# --- TAB 2: DISCOVERY RESPONSE DRAFTER ---
with tab2:
    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("1. Setup Objections & Workflow")
        
        ai_engine = st.radio(
            "Select Inference Model Engine:",
            ["Gemini (Google)", "ChatGPT (OpenAI)"],
            horizontal=True
        )
        
        user_api_key = st.text_input(
            f"Required: Enter your personal {ai_engine} API Key",
            type="password",
            placeholder="Key is not logged or stored"
        )
        
        st.markdown("---")
        
        use_case_styling = st.checkbox("Include Case Styling Box at Top?", value=True)
        
        cause_no = ""
        plaintiff = ""
        defendant = ""
        court = ""
        county = ""
        
        if use_case_styling:
            c1, c2 = st.columns(2)
            with c1:
                cause_no = st.text_input("Cause No.", value="DC-25-23025")
                plaintiff = st.text_input("Plaintiff", value="MIGUEL ORTIZ")
                court = st.text_input("Court", value="160th Judicial District")
            with c2:
                defendant = st.text_input("Defendant", value="GAMALIEL MORALES BAUTISTA")
                county = st.text_input("County", value="Dallas")
            
            st.markdown("---")

        request_type = st.selectbox(
            "Discovery Type", 
            ["Requests for Production (RFPs)", "Interrogatories (ROGs)", "Requests for Admissions (RFAs)"]
        )
        
        selected_objections = st.multiselect(
            "Select Candidate Objections to Screen & Apply:",
            list(TAXONOMY_OBJECTIONS.keys())
        )
        
        incoming_request = st.text_area(
            "Paste Defendant's Exact Question",
            height=130,
            placeholder="e.g., Interrogatory No. 1: Please state your full name..."
        )
        
        factual_basis = st.text_area(
            "Enter Factual Answer/Response details",
            height=100,
            placeholder="What actually occurred or what do we have?"
        )
        
        run_button = st.button("Generate Final Discovery Response", type="primary")

    with col2:
        st.subheader("2. Live Response Preview")
        
        if run_button:
            if not incoming_request:
                st.warning("Please paste the defendant's request first.")
            elif not user_api_key:
                st.error(f"Please provide your personal {ai_engine} API Key.")
            else:
                objection_text = "\n".join([f"- {TAXONOMY_OBJECTIONS[obj]}" for obj in selected_objections])
                
                prompt = f"""You are a meticulous legal discovery drafting engine. Your job is to output a direct legal response.

INCOMING REQUEST TYPE: {request_type}
DEFENDANT'S REQUEST:
"{incoming_request}"

PLAINTIFF'S FACTUAL RESPONSE:
"{factual_basis if factual_basis else "Plaintiff has provided no additional information at this time."}"

THE SELECTED CANDIDATE OBJECTIONS TO APPLY:
{objection_text if objection_text else "None."}

STRICT RESPONSE RULES:
1. Avoid general opening context statements or conversation. Do not say "Here is your response".
2. Do not use Markdown formatting characters like '#' or '*' in your output unless wrapping a heading.
3. First, type out the exact text of the relevant objections applied.
4. Directly after the objections, output: "**Subject to and without waiving the foregoing, Plaintiff responds as follows:**"
5. Add the factual response text.
"""
                
                output_text = ""
                
                if ai_engine == "Gemini (Google)":
                    with st.spinner("Processing objections via Gemini..."):
                        try:
                            genai.configure(api_key=user_api_key)
                            model = genai.GenerativeModel("gemini-2.5-flash")
                            response = model.generate_content(prompt)
                            output_text = response.text
                        except Exception as e:
                            st.error(f"Error calling Gemini AI: {e}")
                                
                elif ai_engine == "ChatGPT (OpenAI)":
                    with st.spinner("Processing objections via ChatGPT..."):
                        try:
                            client = OpenAI(api_key=user_api_key)
                            response = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[{"role": "user", "content": prompt}]
                            )
                            output_text = response.choices[0].message.content
                        except Exception as e:
                            st.error(f"Error calling ChatGPT AI: {e}")
                
                if output_text:
                    st.success("Draft Generated Successfully")
                    st.session_state["last_responder_output"] = output_text

        if "last_responder_output" in st.session_state:
            output_text = st.session_state["last_responder_output"]
            
            st.markdown("### Export Word Document")
            
            case_info = {
                "cause_no": cause_no,
                "plaintiff": plaintiff,
                "defendant": defendant,
                "court": court,
                "county": county
            } if use_case_styling else None
            
            docx_data = generate_response_docx(
                content_text=output_text, 
                include_styling_box=use_case_styling, 
                style_details=case_info
            )
            
            st.download_button(
                label="📥 Download Word File (.docx)",
                data=docx_data,
                file_name="Compliant_Discovery_Response.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            st.markdown("---")
            st.markdown(output_text)
