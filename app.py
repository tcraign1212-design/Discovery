import streamlit as st
import io
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
from openai import OpenAI
import anthropic
import os

# ──────────────────────────────────────────────
# 1. PAGE CONFIGURATION
# ──────────────────────────────────────────────
st.set_page_config(page_title="Case Intelligence Brief Generator", layout="wide")

st.title("Legal Utility: Case Intelligence Brief Generator")
st.markdown("*Strategic case framing for pre-litigation intake and active litigation workup*")
st.markdown("---")

# ──────────────────────────────────────────────
# 2. SECURE API KEY RESOLUTION
# ──────────────────────────────────────────────
env_gemini_key    = st.secrets.get("GEMINI_API_KEY")    or os.environ.get("GEMINI_API_KEY")
env_openai_key    = st.secrets.get("OPENAI_API_KEY")    or os.environ.get("OPENAI_API_KEY")
env_anthropic_key = st.secrets.get("ANTHROPIC_API_KEY") or os.environ.get("ANTHROPIC_API_KEY")


# ──────────────────────────────────────────────
# 3. DOCUMENT GENERATION
# ──────────────────────────────────────────────
def generate_brief_docx(brief_text: str, doc_title: str) -> io.BytesIO:
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(12)

    # Document title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(doc_title)
    run_title.font.bold = True
    run_title.font.size = Pt(13)

    # Divider
    doc.add_paragraph("_" * 70)
    doc.add_paragraph()

    # Parse and render markdown-style output
    lines = brief_text.split("\n")
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()

        if cleaned.startswith("### "):
            run = p.add_run(cleaned.replace("### ", "").strip())
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.underline = True

        elif cleaned.startswith("## "):
            run = p.add_run(cleaned.replace("## ", "").strip())
            run.font.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(10)

        elif cleaned.startswith("# "):
            run = p.add_run(cleaned.replace("# ", "").strip())
            run.font.bold = True
            run.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(12)

        elif cleaned.startswith(("- ", "* ")):
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(cleaned[2:].strip())

        elif cleaned.startswith(("**", "__")) and cleaned.endswith(("**", "__")):
            run = p.add_run(cleaned.strip("*_"))
            run.font.bold = True

        else:
            run = p.add_run(cleaned)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ──────────────────────────────────────────────
# 4. PROMPT BUILDER
# ──────────────────────────────────────────────
def build_prompt(
    case_stage: str,
    case_type: str,
    discovery_level: str,
    date_of_incident: str,
    sol_date: str,
    case_summary: str,
    government_entity: bool,
    commercial_vehicle: bool,
) -> str:

    gov_flag = (
        "\n   - GOVERNMENT ENTITY FLAG: User has identified government entity involvement. "
        "Rigorously apply TTCA / sovereign immunity analysis. Flag all notice deadlines precisely."
        if government_entity else ""
    )

    cv_flag = (
        "\n   - COMMERCIAL VEHICLE FLAG: Apply FMCSR analysis. Flag Hours of Service, "
        "qualification file, and black box / EDR preservation obligations."
        if commercial_vehicle else ""
    )

    shared_params = f"""
CASE PARAMETERS:
- Framework: {case_type}
- Date of Incident: {date_of_incident if date_of_incident else "Not provided"}
- SOL Date: {sol_date if sol_date else "Not provided"}
- Government Entity Involvement: {"YES" if government_entity else "No"}
- Commercial Vehicle Involvement: {"YES" if commercial_vehicle else "No"}

CASE SUMMARY:
\"\"\"{case_summary}\"\"\"

RISK FLAGS:{gov_flag}{cv_flag if cv_flag else " None flagged by user."}
"""

    if case_stage == "Pre-Litigation":
        return f"""
You are a Texas personal injury litigation strategist conducting a pre-suit intake review.
Your output is a Case Intelligence Brief — a structured strategic framework for early case workup.
Do NOT draft discovery requests. This matter has not been filed.

{shared_params}

DIRECTIVES:
1. Flag government entity involvement with precision: identify the specific agency, applicable
   immunity doctrine, and the exact notice deadline (TTCA 6-month notice, home-rule municipality
   rules, EMS/fire sovereign immunity, etc.).
2. Identify FOIA / Texas Public Information Act targets: which agencies hold relevant records,
   estimated record categories, and any applicable 10-business-day response deadlines.
3. Identify spoliation risk and what preservation demand letters should issue immediately.
4. Do not draft discovery requests. This is a pre-suit tool.
5. Be precise on dates and deadlines. If the SOL date is provided, flag any approaching windows.
6. Use these exact section headers — no others:

## 1. Chronology & Case Metrics
(Incident date, SOL expiration, days remaining, key timeline milestones)

## 2. Liability Theory & Exposure Analysis
(Specific breach identified | Applicable legal standard | Exposure range if determinable)

## 3. Intake Risk Flags
(Government entities | Notice requirements | SOL traps | Sovereign immunity hurdles)

## 4. Proof Gap Analysis
(What evidence currently exists | What is missing | Who holds it | How to obtain it)

## 5. Defense Anticipation
(Most likely affirmative defenses | Anticipated contention points | Comparative fault exposure)

## 6. Pre-Suit Action Items
(Ordered by urgency: preservation demands, FOIA/PIA requests, notice letters, workup sequence,
recommended records to obtain before filing)
"""

    else:  # Active Litigation
        return f"""
You are a Texas personal injury litigation strategist conducting a case workup review for
an active suit. Your output is a Case Intelligence Brief — a Discovery Blueprint and
strategic framework. Do NOT draft actual discovery requests. Generate the framework
a paralegal will use to draft precise, objection-resistant requests in a separate workflow.

{shared_params}
- Discovery Level: {discovery_level}

DIRECTIVES:
1. Do not include items covered by TRCP Rule 194.2 Required Disclosures.
2. Do not draft actual Interrogatories, RFPs, or RFAs. Identify the targets, the legal
   basis for each target, and why each target matters to the theory of the case.
3. Organize the Discovery Blueprint by request type so a paralegal can draft directly from it.
4. Flag spoliation risk, expert needs, and any third-party liability exposure.
5. Be precise — tailor every point to the specific facts in the case summary.
6. Use these exact section headers — no others:

## 1. Chronology & Case Metrics
(Incident date, filing date if known, discovery deadline, SOL, key milestones)

## 2. Liability Theory & Exposure Analysis
(Specific breach | Legal standard | Comparative fault exposure | Damages overview)

## 3. Proof Gap Analysis
(What evidence exists | What is missing | Who holds it | Priority of acquisition)

## 4. Defense Anticipation & Contention Points
(Affirmative defenses likely to be pled | How to pin each down in discovery)

## 5. Discovery Blueprint
Organized by request type — identify the target and its strategic purpose:
  ### Requests for Admission
  (Facts to establish or liability points to lock down — categorized by topic)
  ### Requests for Production
  (Documents to compel — categorized by topic, with custodian identified where possible)
  ### Interrogatories
  (Information to pin down, especially re: defenses, prior incidents, and causation)

## 6. Strategic Flags
(Expert witness needs | Spoliation risk | Third-party liability | Insurance coverage issues |
 Any FMCSR / TTCA / regulatory overlay requiring specialized discovery)
"""


# ──────────────────────────────────────────────
# 5. MAIN LAYOUT
# ──────────────────────────────────────────────
col1, col2 = st.columns([2, 3])

with col1:
    st.subheader("1. Case Details")

    # ── Case Stage (drives everything) ──
    case_stage = st.radio(
        "Case Stage:",
        ["Pre-Litigation", "Active Litigation"],
        horizontal=True,
    )

    st.markdown("---")

    # ── Model Engine ──
    ai_engine = st.radio(
        "Select Model Engine:",
        ["Gemini (Google)", "ChatGPT (OpenAI)", "Claude (Anthropic)"],
        horizontal=True,
    )

    active_api_key = ""
    if ai_engine == "Gemini (Google)":
        active_api_key = st.text_input(
            "Gemini API Key", value=env_gemini_key or "", type="password"
        )
    elif ai_engine == "ChatGPT (OpenAI)":
        active_api_key = st.text_input(
            "OpenAI API Key", value=env_openai_key or "", type="password"
        )
    elif ai_engine == "Claude (Anthropic)":
        active_api_key = st.text_input(
            "Anthropic API Key", value=env_anthropic_key or "", type="password"
        )

    st.markdown("---")

    # ── Case Type ──
    case_type = st.selectbox(
        "Case Type / Framework:",
        [
            "Standard Motor Vehicle Accident (MVA)",
            "Commercial Vehicle / Trucking Crash",
            "Premises Liability (Slip/Trip/Fall)",
            "Workplace Injury / Non-Subscriber",
            "Uninsured/Underinsured Motorist (UM/UIM)",
            "Texas Tort Claims Act (TTCA) / Sovereign Immunity",
        ],
    )

    # ── Discovery Level — only shown in litigation mode ──
    if case_stage == "Active Litigation":
        discovery_level = st.radio(
            "Discovery Level (TRCP):",
            ["Level 1 (Expedited, up to $250k)", "Level 2 (Standard)", "Level 3 (Custom/Complex)"],
            index=1,
            horizontal=True,
        )
    else:
        discovery_level = "N/A (Pre-Litigation)"

# ── Risk Flags ──
    st.markdown("**Risk Flags** *(check all that apply)*")
    government_entity = st.checkbox("Government Entity Involved (TTCA / Sovereign Immunity)")
    
    commercial_vehicle = st.checkbox("Commercial Vehicle Involved")

    st.markdown("**Commercial Vehicle Status**")
    commercial_status = st.radio(
        "Does FMCSR / Commercial Regs apply?",
        ["Confirmed Yes", "Confirmed No", "Unsure / Needs Analysis"],
        index=2, # Defaults to "Unsure"
        horizontal=True
)

# This variable then feeds into your build_prompt() function
    if commercial_status in ["Confirmed Yes", "Unsure / Needs Analysis"]:
    # The code will now tell the AI to include the FMCSR Research section
        include_fmcsr_analysis = True 
    else:
        include_fmcsr_analysis = False

    # ── Case Summary ──
    case_summary = st.text_area(
        "Case Summary / Fact Pattern",
        height=160,
        placeholder=(
            "Enter key facts: parties, mechanism of injury, known witnesses, "
            "records obtained, insurance status, any known defenses..."
        ),
    )

    # ── Optional Dates ──
    with st.expander("Dates & Deadlines"):
        c1, c2 = st.columns(2)
        with c1:
            date_of_incident = st.text_input("Date of Incident", placeholder="YYYY-MM-DD")
        with c2:
            sol_date = st.text_input("SOL Expiration Date", placeholder="YYYY-MM-DD")

    run_brief = st.button(
        "Generate Case Intelligence Brief", type="primary", use_container_width=True
    )


# ──────────────────────────────────────────────
# 6. OUTPUT PANEL
# ──────────────────────────────────────────────
with col2:
    st.subheader("2. Case Intelligence Brief")

if run_brief:
        if not case_summary.strip():
            st.warning("Please provide a case summary before generating.")
        elif not active_api_key:
            st.error(f"Please enter an API key for {ai_engine}.")
        else:
            prompt = build_prompt(
                case_stage=case_stage,
                case_type=case_type,
                discovery_level=discovery_level,
                date_of_incident=date_of_incident,
                sol_date=sol_date,
                case_summary=case_summary,
                government_entity=government_entity,
                commercial_vehicle=commercial_vehicle,
            )

            output_text = ""

            if ai_engine == "Gemini (Google)":
                with st.spinner("Analyzing via Gemini..."):
                    try:
                        genai.configure(api_key=active_api_key)
                        model = genai.GenerativeModel("gemini-2.5-flash")
                        response = model.generate_content(prompt)
                        output_text = response.text
                    except Exception as e:
                        st.error(f"Gemini error: {e}")

            elif ai_engine == "ChatGPT (OpenAI)":
                with st.spinner("Analyzing via OpenAI..."):
                    try:
                        client = OpenAI(api_key=active_api_key)
                        response = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[{"role": "user", "content": prompt}],
                        )
                        output_text = response.choices[0].message.content
                    except Exception as e:
                        st.error(f"OpenAI error: {e}")

            elif ai_engine == "Claude (Anthropic)":
                with st.spinner("Analyzing via Claude..."):
                    try:
                        client = anthropic.Anthropic(api_key=active_api_key)
                        response = client.messages.create(
                            model="claude-opus-4-5-20251101",
                            max_tokens=4000,
                            messages=[{"role": "user", "content": prompt}],
                        )
                        output_text = response.content[0].text
                    except Exception as e:
                        st.error(f"Claude error: {e}")

            if output_text:
                st.session_state["brief_content"] = output_text
                st.session_state["brief_stage"]   = case_stage

    # ── Editable output area ──
    if "brief_content" in st.session_state:
        stage_label = st.session_state.get("brief_stage", "")
        doc_title = (
            "PRE-LITIGATION CASE INTELLIGENCE BRIEF"
            if stage_label == "Pre-Litigation"
            else "LITIGATION CASE INTELLIGENCE BRIEF"
        )

        st.info(
            "Review and edit the brief below. When satisfied, export to Word. "
            "Use Section 4 or 5 output as your prompt prefix when drafting discovery in Claude or Midpage."
        )

        edited_text = st.text_area(
            "Edit before exporting:",
            value=st.session_state["brief_content"],
            height=500,
        )
        st.session_state["brief_content"] = edited_text

        st.markdown("### Export")
        docx_data = generate_brief_docx(st.session_state["brief_content"], doc_title)
        st.download_button(
            label="📥 Download Case Intelligence Brief (.docx)",
            data=docx_data,
            file_name=f"{'Pre_Lit' if stage_label == 'Pre-Litigation' else 'Litigation'}_Case_Intelligence_Brief.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.markdown("---")
        st.caption(
            "**Next step:** Copy the Proof Gap Analysis and Discovery Blueprint sections into "
            "Claude, Midpage, or Gemini as a structured prompt prefix. Draft discovery requests "
            "against that framework — not from raw facts."
        )
